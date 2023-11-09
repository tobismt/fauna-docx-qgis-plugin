import pandas as pd
import docx
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt
from qgis.core import QgsProject
import os

class roteListePflanzen:
    def __init__(self, pflanzen_layer_name, outpath):
        self.pflanzen_layer = QgsProject.instance().mapLayersByName(pflanzen_layer_name)[0]
        self.outpath = outpath
        self.basepath = os.path.dirname(os.path.realpath(__file__))
        self.LUT = pd.read_csv(f"{self.basepath}/flora.csv", sep='|')
        self.doc = docx.Document()
        
        self.get_arten_list(self.pflanzen_layer)
        self.create_df()
        self.add_header()
        self.df_to_word()
        self.color_cells(self.doc.tables[0])
        self.center_text()
        self.save()
        
    def add_header(self):
        header = self.doc.sections[0].header
        header.paragraphs[0].text = "Rote-Liste Pflanzen im Untersuchungsgebiet\n"
        header.paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        style = self.doc.styles['Heading 1']
        font = style.font
        font.size = Pt(16)
        
        header.paragraphs[0].style = self.doc.styles['Heading 1']
        
    def get_arten_list(self, lyr):
        cols = [f.name() for f in lyr.fields()]
        datagen = ([f[col] for col in cols] for f in lyr.getFeatures())

        df = pd.DataFrame.from_records(data=datagen, columns=cols)
        
        list = df.name.unique()
        self.list = list


    def create_df(self):
        df = pd.DataFrame(self.list, columns = ['Name'])
        merge = df.merge(self.LUT, how='left', left_on='Name', right_on='Name')
        merge = merge[['Name', 'Deutscher Name', 'Aktuelle Bestandssituation', 'Langfristiger Bestandstrend']]
        merge = merge.fillna('')
        merge = merge.sort_values('Name')
        
        self.df = merge

    
    def df_to_word(self):
        t = self.doc.add_table(self.df.shape[0]+1, self.df.shape[1])
        for j in range(self.df.shape[-1]):
            t.cell(0,j).text = self.df.columns[j]

        for i in range(self.df.shape[0]):
            for j in range(self.df.shape[-1]):
                t.cell(i+1,j).text = str(self.df.values[i,j])

    def color_cells(self, table):
        num = 0
        elems = []
        for row in table.rows:
            for cell in row.cells:
                if 'Mäßig häufig' == cell.text:
                    elems.append(parse_xml(r'<w:shd {} w:fill="ddf220"/>'.format(nsdecls('w'))))
                    cell._tc.get_or_add_tcPr().append(elems[num])
                    num += 1
                if 'Sehr häufig' == cell.text:
                    elems.append(parse_xml(r'<w:shd {} w:fill="2cb827"/>'.format(nsdecls('w'))))
                    cell._tc.get_or_add_tcPr().append(elems[num])
                    num += 1
                if 'Sehr selten' == cell.text:
                    elems.append(parse_xml(r'<w:shd {} w:fill="e81010"/>'.format(nsdecls('w'))))
                    cell._tc.get_or_add_tcPr().append(elems[num])
                    num += 1
                if 'Extrem selten' == cell.text:
                    elems.append(parse_xml(r'<w:shd {} w:fill="da097e"/>'.format(nsdecls('w'))))
                    cell._tc.get_or_add_tcPr().append(elems[num])
                    num += 1
                if 'Selten' == cell.text:
                    elems.append(parse_xml(r'<w:shd {} w:fill="f18c18"/>'.format(nsdecls('w'))))
                    cell._tc.get_or_add_tcPr().append(elems[num])
                    num += 1
                if 'Ausgestorben oder verschollen' == cell.text:
                    elems.append(parse_xml(r'<w:shd {} w:fill="c900ff"/>'.format(nsdecls('w'))))
                    cell._tc.get_or_add_tcPr().append(elems[num])
                    num += 1
                if 'Häufig' == cell.text:
                    elems.append(parse_xml(r'<w:shd {} w:fill="9edd23"/>'.format(nsdecls('w'))))
                    cell._tc.get_or_add_tcPr().append(elems[num])
                    num += 1
                    
                    
    def center_text(self):
        for row in self.doc.tables[0].rows:
            if row.cells[1].text == 'Deutscher Name':
                for cell in row.cells:
                    paragraphs = cell.paragraphs
                    for paragraph in paragraphs:
                        for run in paragraph.runs:
                            font = run.font
                            font.size= Pt(12)
                            font.bold = True
            for cell in row.cells:
                    paragraphs = cell.paragraphs
                    for paragraph in paragraphs:
                        paragraph.paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
            
        
    def save(self):
        self.doc.save(self.outpath)
        