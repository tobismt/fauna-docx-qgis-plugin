import pandas as pd
import docx
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt, Cm
import os

class redListFauna:
    def __init__(self, fauna_layer, field, outpath):
        """
        Constructor for the redListFauna class.

        Parameters:
        - fauna_layer: QgsVectorLayer, the vector layer containing fauna data.
        - field: str, the field name to be used in the table.
        - outpath: str, the output path for the generated Word document.
        """
        self.fauna_layer = fauna_layer
        self.outpath = outpath
        self.field = field
        self.basepath = os.path.dirname(os.path.realpath(__file__))
        self.LUT = pd.read_csv(f"{self.basepath}/fauna.csv", sep='|')  # Look-up table for fauna data
        self.legend = pd.read_csv(f"{self.basepath}/legend.csv", sep='|')  # Legend for table colors
        self.doc = docx.Document()

        self.get_arten_list(self.fauna_layer)  # Retrieve unique fauna names from the layer
        self.create_df()  # Create a DataFrame with relevant fauna data
        self.add_header()  # Add document header
        self.df_to_word()  # Convert DataFrame to Word table
        self.color_cells(self.doc.tables[0])  # Apply color to cells based on values
        self.center_text()  # Center-align text in the table
        self.create_legend()  # Add legend to the document
        self.save()  # Save the document

    def add_header(self):
        """
        Add header to the document.
        """
        header = self.doc.sections[0].header
        header.paragraphs[0].text = "Rote-Liste Fauna im Untersuchungsgebiet\n"
        header.paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

        style = self.doc.styles['Heading 1']
        font = style.font
        font.size = Pt(16)

        header.paragraphs[0].style = self.doc.styles['Heading 1']

    def get_arten_list(self, lyr):
        """
        Extract unique fauna names from the provided vector layer.

        Parameters:
        - lyr: QgsVectorLayer, vector layer containing fauna data.
        """
        cols = [f.name() for f in lyr.fields()]
        datagen = ([f[col] for col in cols] for f in lyr.getFeatures())

        df = pd.DataFrame.from_records(data=datagen, columns=cols)

        list = df.Name.unique()
        self.list = list

    def create_df(self):
        """
        Create DataFrame with relevant fauna data for the specified field.
        """
        df = pd.DataFrame(self.list, columns=[self.field])
        merge = df.merge(self.LUT, how='left', left_on=self.field, right_on='Name')
        merge = merge[['Name', 'Deutscher Name', 'aktuelle Bestandssituation', 'kurzfristiger Bestandstrend',
                       'langfristiger Bestandstrend', 'RL Kat.']]
        merge.columns = [i.title() for i in merge.columns]
        merge = merge.fillna('-')
        merge = merge.sort_values('Name')

        self.df = merge

    def df_to_word(self):
        """
        Convert DataFrame to Word table.
        """
        t = self.doc.add_table(self.df.shape[0] + 1, self.df.shape[1], style="Table Grid")

        # Transfer DataFrame content to docx table
        for j in range(self.df.shape[-1]):
            t.cell(0, j).text = self.df.columns[j]

        for i in range(self.df.shape[0]):
            for j in range(self.df.shape[-1]):
                t.cell(i + 1, j).text = str(self.df.values[i, j])

        # Set column widths
        for i, col in enumerate(t.columns):
            if i == 0 or i == 1:
                col.width = Cm(4)
            else:
                col.width = Cm(2.5)

    def color_cells(self, table):
        """
        Apply color to cells based on values in the table.

        Parameters:
        - table: docx.table.Table, the table to be colored.
        """
        num = 0
        elems = []
        for row in table.rows:
            for cell in row.cells:
                # Define cell colors based on cell values
                if '0' == cell.text:
                    elems.append(parse_xml(r'<w:shd {} w:fill="#3ec902"/>'.format(nsdecls('w'))))
                    cell._tc.get_or_add_tcPr().append(elems[num])
                    num += 1
                if '1' == cell.text:
                    elems.append(parse_xml(r'<w:shd {} w:fill="#80c902"/>'.format(nsdecls('w'))))
                    cell._tc.get_or_add_tcPr().append(elems[num])
                    num += 1
                if '2' == cell.text:
                    elems.append(parse_xml(r'<w:shd {} w:fill="#c9b202"/>'.format(nsdecls('w'))))
                    cell._tc.get_or_add_tcPr().append(elems[num])
                    num += 1
                if '3' == cell.text:
                    elems.append(parse_xml(r'<w:shd {} w:fill="#a30202"/>'.format(nsdecls('w'))))
                    cell._tc.get_or_add_tcPr().append(elems[num])
                    num += 1
                if 'G' == cell.text:
                    elems.append(parse_xml(r'<w:shd {} w:fill="#fa7000"/>'.format(nsdecls('w'))))
                    cell._tc.get_or_add_tcPr().append(elems[num])
                    num += 1
                if 'R' == cell.text:
                    elems.append(parse_xml(r'<w:shd {} w:fill="#b300fa"/>'.format(nsdecls('w'))))
                    cell._tc.get_or_add_tcPr().append(elems[num])
                    num += 1
                if 'V' == cell.text:
                    elems.append(parse_xml(r'<w:shd {} w:fill="#2302c9"/>'.format(nsdecls('w'))))
                    cell._tc.get_or_add_tcPr().append(elems[num])
                    num += 1
                if 'D' == cell.text:
                    elems.append(parse_xml(r'<w:shd {} w:fill="#acacad"/>'.format(nsdecls('w'))))
                    cell._tc.get_or_add_tcPr().append(elems[num])
                    num += 1
                if '*' == cell.text:
                    elems.append(parse_xml(r'<w:shd {} w:fill="#acadad"/>'.format(nsdecls('w'))))
                    cell._tc.get_or_add_tcPr().append(elems[num])
                    num += 1
                if '♦' == cell.text:
                    elems.append(parse_xml(r'<w:shd {} w:fill="9edd23"/>'.format(nsdecls('w'))))
                    cell._tc.get_or_add_tcPr().append(elems[num])
                    num += 1

    def create_legend(self):
        """
        Add legend to the document.
        """
        self.doc.add_paragraph('')
        self.legend.fillna('', inplace=True)
        t = self.doc.add_table(self.legend.shape[0] + 1, self.legend.shape[1], style="Table Grid")

        # Transfer DataFra,e content to docx table
        for j in range(self.legend.shape[-1]):
            t.cell(0, j).text = self.legend.columns[j]

        for i in range(self.legend.shape[0]):
            for j in range(self.legend.shape[-1]):
                if '-' not in str(self.legend.values[i, j]):
                    t.cell(i + 1, j).text = str(self.legend.values[i, j])

        t.alignment = WD_TABLE_ALIGNMENT.CENTER

        t.autofit = False
        t.allow_autofit = False
        for col in t.columns:
            col.width = Cm(1.75)

        for i, row in enumerate(t.rows):
            if i == 0:
                # Merge cells and set legend labels
                row.cells[0].merge(row.cells[1])
                row.cells[0].text = 'Rote Liste Status'
                row.cells[2].merge(row.cells[3])
                row.cells[3].text = 'Aktuelle Bestandssituation'
                row.cells[4].merge(row.cells[5])
                row.cells[5].text = 'Bestandstrend langfristig'
                row.cells[6].merge(row.cells[7])
                row.cells[7].text = 'Bestandstrend kurzfristig'
            # Format legend text
            for cell in row.cells:
                if '-' in cell.text:
                    cell.text = ''
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    for run in paragraph.runs:
                        font = run.font
                        if row.cells[0].text == 'Rote Liste Status':
                            font.bold = True
                        font.size = Pt(6)
                    paragraph.paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

    def center_text(self):
        """
        Center-align text in the main table.
        """
        for row in self.doc.tables[0].rows:
            if row.cells[1].text == 'Deutscher Name':
                # Format header row
                for cell in row.cells:
                    paragraphs = cell.paragraphs
                    for paragraph in paragraphs:
                        for run in paragraph.runs:
                            font = run.font
                            font.size = Pt(12)
                            font.bold = True
            for cell in row.cells:
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    paragraph.paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

    def save(self):
        """
        Save document.
        """
        self.doc.save(self.outpath)